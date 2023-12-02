# This module contains a function to extract data from a docx file
import docx
import re
import openpyxl
def extract_data(
        docx_filename, 
        excel_file_name, 
        answer_table: bool, 
        is_underline: bool,
        is_italic: bool):

    # Create a document object from the docx file
    doc = docx.Document(docx_filename)
    wb = openpyxl.Workbook()
    wb = openpyxl.load_workbook(excel_file_name)
    ws = wb.active
    # Initialize an empty list to store the extracted data
    data = []
    # List convert A B C D to number
    answer_index = {"A": 1, "B": 2, "C": 3, "D": 4}
    row = 2
    # Loop through the paragraphs in the document and append their text to the list
    for para in doc.paragraphs:
        text = para.text.strip()
        data.append(text)
        if is_underline:
            for run in para.runs:
                match = re.match(r"(\d+)[.:]\s+(.*)$", text)
                if match:
                    row = int(match.group(1)) + 2

                if run.underline and r"[A-D]\. .*":
                    clean_data_ans = str(run.text).replace(".", "")
                    ws[f"G{row}"] = answer_index[clean_data_ans]   
        
        if is_italic:
            for run in para.runs:
                match = re.match(r"(\d+)[.:]\s+(.*)$", text)
                if match:
                    row = int(match.group(1)) + 2
                else:
                    row += 1
                if run.italic and re.match(r"[A-D]\. .*", run.text):
                    clean_data_ans = str(run.text).split(".")
                    # print(row, clean_data_ans[0].strip())
                    ws[f"G{row}"] = answer_index.get(clean_data_ans[0].strip(), None)    
    if answer_table:
        # print(answer_table)
        table = doc.tables[0]
        # print(table)
        values = []
        for row in table.rows:
            values.append([cell.text.strip() for cell in row.cells])
        # Print the values of the table
        for row, items in enumerate(values):
            for column, cell_value in enumerate(items):
                    # number, word = cell_value.split(' - ')
                    if "–" in cell_value:
                        number, word = cell_value.split('–')
                    # Split using "-" (hyphen)
                    elif "-" in cell_value:
                        number, word = cell_value.split('-')
                    else:
                        continue
                    number = int(number.strip())
                    word = word.strip()
                    print(number, word)
                    ws[f"G{number + 2}"] = answer_index[word]
    
    data = "\n\n".join(data)
        
    # Split the data by double newline characters and assign it to the data_list variable
    data_list = data.split("\n\n")
    wb.save(excel_file_name)
    # Return the data_list as the output of this function
    return data_list
def extra_options(data_list, index):
    string = ''
    match = re.match(r"([A-D]\.) (.*)", data_list[index])
    if match:
        string = match.group(2)
    else:
        print("WRONG FORMAT: ", data_list[index])
        suggestion = re.sub(r"([A-D])\. *", r"\1. ", data_list[index])
        print("PLEASE FIX:", data_list[index], "to", suggestion, "in your docx file")
    return string
def extra_questions(data_list, index):
    string = ''
    match = re.match(r"(\d+)[.:]\s+(.*)$", data_list[index])
    if match:
        string = match.group(2)
    elif re.match(r""):
        pass
    return string
