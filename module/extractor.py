import docx
import re
import openpyxl

def EXTRACT_DATA(docx_filename, excel_file_name, answer_table: bool, is_underline: bool, is_italic: bool):
    doc = docx.Document(docx_filename)
    wb = openpyxl.load_workbook(excel_file_name)
    ws = wb.active

    DATA = []
    ANSWER_INDEX = {"A": 1, "B": 2, "C": 3, "D": 4}
    NUM_QUEST = 0
    ROW = 1

    def PROCESS_TEXT(para):
        nonlocal ROW, NUM_QUEST
        text = para.text.strip()
        DATA.append(text)

        # Process data-raw
        if re.match(r"(?<!\b[A-D]\.\s)(?<!-\W)\b(?:Câu\s*)?(\d+)[.:]", text):
            NUM_QUEST += 1
            replace_quest = re.sub(r"(?<!\b[A-D]\.\s)(?<!-\W)\b(?:Câu\s*)?(\d+)[.:]", f"Câu {NUM_QUEST}.", text)
            para.text = replace_quest
        match = re.match(r"Câu (\d+)[.:]\s+(.*)$", text)
        if match:
            ROW = int(match.group(1)) + 2
       
        for run in para.runs:
                if (is_underline and run.underline) or (is_italic and run.italic):
                    clean_data_ans = run.text.replace(".", "").strip()
                    if re.match(r"[A-D]", clean_data_ans):
                        ws[f'G{ROW}'] = ANSWER_INDEX.get(clean_data_ans[0], None)


    for para in doc.paragraphs:
        PROCESS_TEXT(para)
    doc.save(docx_filename)
    if answer_table:
        table = doc.tables[0]
        values = [[cell.text.strip() for cell in row.cells] for row in table.rows]

        for row, items in enumerate(values):
            for cell_value in items:
                if "-" in cell_value:
                    number, word = cell_value.split('-')
                else:
                    continue

                number = int(number.strip())
                word = word.strip()
                ws[f'G{number + 2}'] = ANSWER_INDEX.get(word, None)

    data = "\n\n".join(DATA)
    DATA_LIST = data.split("\n\n")
    wb.save(excel_file_name)
    return DATA_LIST

def EXTRA_OPTIONS(DATA_LIST, index):
    match = re.match(r"([A-D]\.) (.*)", DATA_LIST[index])
    if match:
        return match.group(2)
    else:
        suggestion = re.sub(r"([A-D])\. *", r"\1. ", DATA_LIST[index])
        return suggestion

def EXTRA_QUESTIONS(DATA_LIST, index):
    match = re.match(r"Câu (\d+)[.:]\s+(.*)", DATA_LIST[index])
    if match:
        return match.group(2)
    return ''
