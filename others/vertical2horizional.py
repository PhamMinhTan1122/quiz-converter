import os
import docx

def process_word_table(input_file_path):
    num = []
    char = []

    doc = docx.Document(input_file_path)

    for table in doc.tables:
        for row in table.rows:
            # Assuming there are at least two cells in each row
            col1_value = row.cells[0].text
            col2_value = row.cells[1].text

            num.append(col1_value)
            char.append(col2_value)

    return num, char

def convert_sequence_to_table(numbers, characters, output_file_path):
    document = docx.Document()
    table = document.add_table(rows=0, cols=10)
    table.style = 'Table Grid'

    for i in range(0, len(numbers), 10):
        row_cells = table.add_row().cells
        for j in range(10):
            index = i + j
            if index < len(numbers):
                row_cells[j].text = f"{numbers[index]} - {characters[index]}"

    document.save(output_file_path)

if __name__ == "__main__":
    # Replace 'your_input_word_file.docx' with the path to your input Word file
    # Replace 'output_word_file.docx' with the desired output Word file path
    input_word_file_path = "table_1.docx"
    output_word_file_path = "D:/ProjectPython/quiz-converter/others/output_word_file.docx"
    
    num, char = process_word_table(input_word_file_path)
    convert_sequence_to_table(num, char, output_word_file_path)
